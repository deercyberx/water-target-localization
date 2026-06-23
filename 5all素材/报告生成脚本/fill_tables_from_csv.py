# -*- coding: utf-8 -*-
"""从CSV原始数据重新计算并填充v9_final的12张表"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import pandas as pd
import numpy as np
import json
import yaml
from docx import Document

# ============================================================
# 1. 从CSV计算所有指标
# ============================================================
df_high = pd.read_csv('素材/实验数据/results_HIGH.csv')
df_low = pd.read_csv('素材/实验数据/results_LOW.csv')
df_top1 = pd.read_csv('素材/实验数据/results_Top1.csv')
df_most = pd.read_csv('素材/实验数据/results_MostInliers.csv')
df_sift = pd.read_csv('素材/实验数据/results_SIFT.csv')

def calc(errors):
    e = np.array(errors); e = e[np.isfinite(e)]
    return {
        'n': len(e),
        'a5': round(100*np.mean(e<=5),1), 'a10': round(100*np.mean(e<=10),1), 'a20': round(100*np.mean(e<=20),1),
        'mean': round(np.mean(e),1), 'median': round(np.median(e),1)
    }

m_high = calc(df_high['pred_error'])
m_low = calc(df_low['pred_error'])

scenes = {}
for s in df_high['scene'].unique():
    scenes[s] = calc(df_high[df_high['scene']==s]['pred_error'])

# 俯仰角分桶（修正版，include_lowest=True）
bins = [-90, -70, -50, -20, 0]
labels = ['70-90°', '50-70°', '20-50°', '0-20°']
df_high['pbin'] = pd.cut(df_high['pitch'], bins=bins, labels=labels, right=True, include_lowest=True)
pitch_data = {}
for l in labels:
    sub = df_high[df_high['pbin']==l]
    if len(sub)>0: pitch_data[l] = calc(sub['pred_error'])

# 高度分桶
df_high['abin'] = pd.cut(df_high['altitude'], bins=[0,100,200,300], labels=['30-100m','100-200m','200-300m'], right=True, include_lowest=True)
alt_data = {}
for l in ['30-100m','100-200m','200-300m']:
    sub = df_high[df_high['abin']==l]
    if len(sub)>0: alt_data[l] = calc(sub['pred_error'])

m_top1 = calc(df_top1['pred_error'])
m_most = calc(df_most['pred_error'])
m_sift = calc(df_sift['pred_error'])

with open('code/config.yaml') as f: config = yaml.safe_load(f)
with open('code/Data/metadata/QZ_Town.json') as f: meta = json.load(f)

# ============================================================
# 2. 填充v9_final的12张表
# ============================================================
doc = Document('reports/第五章_实验验证与结果分析_poi-tl_v9_final.docx')

def set_cell(table, r, c, text):
    cell = table.rows[r].cells[c]
    for p in cell.paragraphs:
        for run in p.runs: run.text = ''
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = str(text)
    elif cell.paragraphs:
        cell.paragraphs[0].text = str(text)

tables = doc.tables

# --- 表5-1: 数据集属性 ---
t = tables[0]
data1 = [
    ('数据集名称', 'AnyVisLoc'), ('论文来源', 'CVPR 2026 Findings'),
    ('UAV图像数量', f'{len(meta)}张'), ('参考图数量', '196张（124航空+72卫星）'),
    ('机型', '7种DJI'), ('城市/区域', '15城市/25区域/304场景'),
    ('飞行高度', '30m-300m'), ('俯仰角范围', '20°-90°'),
    ('图像分辨率', '1920×1080 ~ 5472×3648'), ('焦距范围', '4.5-28mm'),
    ('航空图分辨率', '0.02-0.35m'), ('卫星图分辨率', '0.13-0.55m'), ('参考图坐标系', 'UTM'),
]
for i, (k,v) in enumerate(data1):
    set_cell(t, i, 0, k); set_cell(t, i, 1, v)

# --- 表5-2: Demo子集 ---
t = tables[1]
rows2 = [
    ('QZ_SongCity','286','result_roi.tif (0.061m/pix)','satellite_roi.tif (0.260m/pix)','GPS+IMU'),
    ('Qingzhou_3_2','59','同上','同上','同上'),
    ('QingZhou_2024','142','同上','同上','同上'),
    ('总计','487','—','—','—'),
]
for i, row in enumerate(rows2):
    for j, val in enumerate(row): set_cell(t, i+1, j, val)

# --- 表5-3: 参考地图 ---
t = tables[2]
rows3 = [
    ('航空正射影像(DOM)','SfM构建的2D正射影像','提供二维纹理锚点','0.061m/pixel','与UAV图像时相一致'),
    ('航空DSM','SfM生成的数字表面模型','二维匹配点升维为三维约束','0.937m/pixel','高程精度高'),
    ('卫星正射影像','Google Earth历史图像','航空图不可用时的替代','0.260m/pixel','存在时相差异'),
    ('卫星DSM','ALOS 30m DSM','卫星图条件下的三维约束','30m/pixel','高程精度低'),
]
for i, row in enumerate(rows3):
    for j, val in enumerate(row): set_cell(t, i+1, j, val)

# --- 表5-4: 实验平台 ---
t = tables[3]
rows4 = [
    ('操作系统','Windows 11'),('CPU','Intel Core i9-10920X @ 3.50GHz'),
    ('GPU','NVIDIA GeForce RTX 3080'),('显存','10GB'),('内存','128GB'),
    ('Python','3.9.13'),('深度学习框架','PyTorch 2.2.1+cu121'),('CUDA','12.1'),
    ('检索模型','CAMP (ConvNeXt-Base)'),('匹配模型','RoMa (DINOv2-ViT-L)'),
]
for i, (k,v) in enumerate(rows4):
    set_cell(t, i+1, 0, k); set_cell(t, i+1, 1, v)

# --- 表5-5: 关键参数 (8行x4列: 1表头+7数据行) ---
t = tables[4]
rows5 = [
    ('图像检索','RETRIEVAL_TOPN',str(config['RETRIEVAL_TOPN']),'候选参考图块数量'),
    ('图像检索','RETRIEVAL_COVER',str(config['RETRIEVAL_COVER']),'检索块重叠率(%)'),
    ('图像检索','BATCH_SIZE',str(config['BATCH_SIZE']),'特征提取批大小'),
    ('图像检索','resize_ratio','0.2','图像缩放比例'),
    ('像素匹配','coarse_res','280 (原560)','RoMa粗分辨率'),
    ('像素匹配','upsample_res','512 (原864)','RoMa上采样分辨率'),
    ('位姿解算','strategy','Topn_opt / Top1 / MostInliers','定位策略'),
]
for i, row in enumerate(rows5):
    for j, val in enumerate(row): set_cell(t, i+1, j, val)

# --- 表5-6: 评价指标 ---
t = tables[5]
rows6 = [
    ('A@5m','定位误差<5m的样本比例','越高越好','所有定位实验','定位精度'),
    ('A@10m','定位误差<10m的样本比例','越高越好','所有定位实验','定位精度'),
    ('A@20m','定位误差<20m的样本比例','越高越好','所有定位实验','定位精度'),
    ('平均误差','所有样本定位误差的算术平均','越低越好','所有定位实验','定位精度'),
    ('中位数误差','所有样本定位误差的中位数','越低越好','所有定位实验','定位精度'),
    ('PDM@K','检索结果与真值的空间重叠度','越高越好','检索质量评价','图像检索'),
    ('重投影误差','2D-3D点对应关系的拟合质量','越低越好','PnP解算质量','位姿解算'),
    ('内点数量','PnP RANSAC内点数','越多越好','策略对比实验','位姿解算'),
    ('推理时间','单帧端到端处理时间','越低越好','性能评估','全链路'),
]
for i, row in enumerate(rows6):
    for j, val in enumerate(row): set_cell(t, i+1, j, val)

# --- 表5-7: 主结果 ---
t = tables[6]
for ci, val in enumerate([str(m_high['n']), f"{m_high['a5']}%", f"{m_high['a10']}%", f"{m_high['a20']}%", f"{m_high['mean']}m", f"{m_high['median']}m"]):
    set_cell(t, 1, ci+1, val)
for ci, val in enumerate([str(m_low['n']), f"{m_low['a5']}%", f"{m_low['a10']}%", f"{m_low['a20']}%", f"{m_low['mean']}m", f"{m_low['median']}m"]):
    set_cell(t, 2, ci+1, val)

# --- 表5-8: 参考图对比 ---
t = tables[7]
for ci, val in enumerate(['0.061m/pixel','0.937m',f"{m_high['a5']}%",f"{m_high['a10']}%",f"{m_high['a20']}%",f"{m_high['mean']}m"]):
    set_cell(t, 1, ci+1, val)
for ci, val in enumerate(['0.260m/pixel','30m',f"{m_low['a5']}%",f"{m_low['a10']}%",f"{m_low['a20']}%",f"{m_low['mean']}m"]):
    set_cell(t, 2, ci+1, val)

# --- 表5-9: 高度与视角 ---
t = tables[8]
row_idx = 1
for label, m in pitch_data.items():
    for ci, val in enumerate(['俯仰角', label, str(m['n']), f"{m['a5']}%", f"{m['a20']}%", f"{m['median']}m"]):
        set_cell(t, row_idx, ci, val)
    row_idx += 1
for label, m in alt_data.items():
    for ci, val in enumerate(['飞行高度', label, str(m['n']),
        f"{m['a5']}%" if not np.isnan(m['a5']) else '—',
        f"{m['a20']}%" if not np.isnan(m['a20']) else '—',
        f"{m['median']}m" if not np.isnan(m['median']) else '—']):
        set_cell(t, row_idx, ci, val)
    row_idx += 1
for s in ['QZ_SongCity','Qingzhou_3_2','QingZhou_2024']:
    m = scenes[s]
    for ci, val in enumerate(['分场景', s, str(m['n']), f"{m['a5']}%", f"{m['a20']}%", f"{m['median']}m"]):
        set_cell(t, row_idx, ci, val)
    row_idx += 1

# --- 表5-10: 策略匹配噪声 (9行x7列: 1表头+8数据行) ---
t = tables[9]
# 策略(3行: 1-3)
for ri, m in enumerate([m_high, m_top1, m_most]):
    for ci, val in enumerate([f"{m['a5']}%", f"{m['a20']}%", f"{m['mean']}m"]):
        set_cell(t, ri+1, ci+2, val)
# 匹配(2行: 4-5)
for ri, m in enumerate([m_high, m_sift]):
    for ci, val in enumerate([f"{m['a5']}%", f"{m['a20']}%", f"{m['mean']}m"]):
        set_cell(t, ri+4, ci+2, val)
# 噪声(3行: 6-8，只放3行yaw数据)
noise_rows = [
    ('Yaw噪声','0°','74.6%','94.2%','—','偏航角噪声影响最大'),
    ('Yaw噪声','30°','70.5%','90.2%','—','标准差30°时A@5m下降4.1pp'),
    ('Yaw噪声','60°','48.9%','70.0%','—','标准差60°时A@5m下降25.7pp'),
]
for ri, row in enumerate(noise_rows):
    for ci, val in enumerate(row):
        set_cell(t, ri+6, ci, val)

# ============================================================
# 3. 保存
# ============================================================
doc.save('reports/第五章_实验验证与结果分析_poi-tl_v9_final.docx')
print('已保存')
print(f'表格数: {len(doc.tables)}, 段落数: {len(doc.paragraphs)}')
print(f'航空图: A@5m={m_high["a5"]}%, 中位数={m_high["median"]}m')
print(f'卫星图: A@5m={m_low["a5"]}%, 中位数={m_low["median"]}m')
print(f'俯仰角: 70-90° n={pitch_data["70-90°"]["n"]}, 50-70° n={pitch_data["50-70°"]["n"]}, 20-50° n={pitch_data["20-50°"]["n"]}')
