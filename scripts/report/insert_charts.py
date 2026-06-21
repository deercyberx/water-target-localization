# -*- coding: utf-8 -*-
"""
插入剩余图表到报告
- matching_comparison.png → 5.8 匹配方法对比
- overall_comparison.png → 5.5 基线结果
- ref_map_comparison.png → 5.6 参考图对比
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

INPUT_FILE = '水面目标精准定位技术报告_终稿_final.docx'
OUTPUT_FILE = '水面目标精准定位技术报告_终稿_v4.docx'

def insert_image_after_paragraph(doc, para_index, image_path, caption, width_inches=5.33):
    """在指定段落后插入图片和图题"""
    # 找到目标段落
    target_para = doc.paragraphs[para_index]

    # 创建新段落用于图片
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 插入图片
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # 将新段落移动到目标段落后面
    target_para._element.addnext(new_para._element)

    # 创建图题段落
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    caption_run.font.size = Pt(10.5)
    caption_run.bold = True

    # 将图题移动到图片后面
    new_para._element.addnext(caption_para._element)

    return caption_para

def find_paragraph_by_text(doc, text, style=None):
    """根据文本内容查找段落"""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            if style is None or p.style.name == style:
                return i
    return -1

def main():
    print(f'读取: {INPUT_FILE}')
    doc = Document(INPUT_FILE)

    print('=' * 50)
    print('插入剩余图表')
    print('=' * 50)

    # 1. 插入 matching_comparison.png 到 5.8 匹配方法对比
    print('\\n[1] 插入匹配方法对比图')
    # 找到"表5-7 不同匹配方法的精度对比"后面
    idx = find_paragraph_by_text(doc, '表5-7')
    if idx == -1:
        idx = find_paragraph_by_text(doc, '不同匹配方法的精度对比')
    if idx >= 0:
        # 找到表格后面的段落
        for i in range(idx, min(idx+10, len(doc.paragraphs))):
            if doc.paragraphs[i].style.name == 'Normal' and doc.paragraphs[i].text.strip():
                idx = i
                break
        insert_image_after_paragraph(
            doc, idx,
            '素材/截图/matching_comparison.png',
            '图5-8 不同匹配方法精度对比'
        )
        print(f'  插入到段落 {idx} 后面')
    else:
        print('  未找到插入位置')

    # 2. 插入 overall_comparison.png 到 5.5 基线结果
    print('\\n[2] 插入复现vs论文对比图')
    # 找到"表5-3 最佳组合基线实验结果"后面
    idx = find_paragraph_by_text(doc, '表5-3')
    if idx == -1:
        idx = find_paragraph_by_text(doc, '最佳组合基线实验结果')
    if idx >= 0:
        # 找到表格后面的段落
        for i in range(idx, min(idx+10, len(doc.paragraphs))):
            if doc.paragraphs[i].style.name == 'Normal' and doc.paragraphs[i].text.strip():
                idx = i
                break
        insert_image_after_paragraph(
            doc, idx,
            '素材/截图/overall_comparison.png',
            '图5-9 复现结果与论文结果对比'
        )
        print(f'  插入到段落 {idx} 后面')
    else:
        print('  未找到插入位置')

    # 3. 插入 ref_map_comparison.png 到 5.6 参考图对比
    print('\\n[3] 插入参考图对比图')
    # 找到"5.6 不同参考地图条件下的定位效果分析"
    idx = find_paragraph_by_text(doc, '5.6')
    if idx == -1:
        idx = find_paragraph_by_text(doc, '不同参考地图条件下的定位效果分析')
    if idx >= 0:
        # 找到该节的第一个Normal段落
        for i in range(idx, min(idx+10, len(doc.paragraphs))):
            if doc.paragraphs[i].style.name == 'Normal' and doc.paragraphs[i].text.strip():
                idx = i
                break
        insert_image_after_paragraph(
            doc, idx,
            '素材/截图/ref_map_comparison.png',
            '图5-10 不同参考地图定位精度对比'
        )
        print(f'  插入到段落 {idx} 后面')
    else:
        print('  未找到插入位置')

    # 保存
    print(f'\\n保存: {OUTPUT_FILE}')
    doc.save(OUTPUT_FILE)
    print('完成!')

    # 统计
    doc2 = Document(OUTPUT_FILE)
    print(f'\\n最终统计:')
    print(f'  段落数: {len(doc2.paragraphs)}')
    print(f'  表格数: {len(doc2.tables)}')
    print(f'  图片数: {len(doc2.inline_shapes)}')

if __name__ == '__main__':
    main()
