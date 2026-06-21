# -*- coding: utf-8 -*-
"""
用 python-docx 生成第五章 docx（直接生成，不依赖poi-tl模板渲染）
复用报告v2脚本的格式设置，读取已生成的图表和表格数据
"""

import sys, os, json, glob
import numpy as np
import pandas as pd
from pathlib import Path

# 添加项目路径
BASE = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(BASE))

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ============================================================
# 工具函数
# ============================================================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def add_table(doc, headers, rows, caption=None):
    """添加表格"""
    if caption:
        p = doc.add_paragraph(caption, style='Heading 3')

    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # 表后空行
    return table

def add_figure(doc, img_path, caption=None):
    """添加图片"""
    if img_path and Path(img_path).exists():
        doc.add_picture(str(img_path), width=Cm(15))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption, style='Heading 3')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_para(doc, text, bold=False):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

# ============================================================
# 主函数
# ============================================================

def generate_chapter5():
    print("生成第五章...")

    doc = Document()

    # 设置页边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)  # 段后6pt

    # 设置标题样式
    for level, (size, name) in enumerate([(16, 'SimHei'), (14, 'SimHei'), (12, 'SimHei')], 1):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = name
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), name)

    # 图表路径
    FIG = BASE / '素材' / '截图'
    TAB = BASE / '素材' / '实验数据'

    # ============================================================
    # 5.1 实验数据与测试场景
    # ============================================================
    doc.add_heading('5.1 实验数据与测试场景', level=1)

    add_para(doc,
        '本章基于CVPR 2026 Findings论文发布的AnyVisLoc数据集的Demo版本（1/25），'
        '对低空多视角UAV视觉定位链路进行系统性测试。测试区域为青州古镇（QZ_Town），'
        '包含3个子场景：QZ_SongCity（286张）、Qingzhou_3_2（59张）、QingZhou_2024（142张），'
        '共计487张无人机图像。'
    )

    add_para(doc,
        '需要说明的是，本章使用的数据主要用于验证低空视觉定位基础链路，'
        '为后续目标级定位提供UAV位姿解算基础，而非直接用于水面人体目标定位实测。'
    )

    # 表5-1
    add_table(doc,
        ['属性', '内容'],
        [
            ['数据集名称', 'AnyVisLoc'],
            ['论文来源', 'CVPR 2026 Findings'],
            ['UAV图像数量', '18,000张'],
            ['参考图数量', '196张（124航空 + 72卫星）'],
            ['机型', '7种DJI（Mavic 2/3/3 Pro, Phantom 3/4/4 RTK, Mini 4 Pro）'],
            ['飞行高度', '30m – 300m'],
            ['俯仰角范围', '20° – 90°'],
            ['数据用途', 'UAV视觉定位链路测试'],
        ],
        caption='表5-1 AnyVisLoc数据集属性'
    )

    # 表5-2
    add_table(doc,
        ['测试区域', '查询图像数', '航空参考图', '卫星参考图', '姿态元数据'],
        [
            ['QZ_SongCity', '286', 'result_roi.tif (0.061m/pix)', 'satellite_roi.tif (0.260m/pix)', 'GPS+IMU'],
            ['Qingzhou_3_2', '59', '同上', '同上', '同上'],
            ['QingZhou_2024', '142', '同上', '同上', '同上'],
            ['总计', '487', '—', '—', '—'],
        ],
        caption='表5-2 Demo测试子集说明'
    )

    add_figure(doc, FIG / 'fig5_1_sample_images.png', '图5-1 测试区域图像样例')

    add_para(doc,
        '数据集使用7款主流大疆无人机采集，飞行高度30m-300m，俯仰角20°-90°，'
        '覆盖多种天气、季节和光照条件。'
    )

    # ============================================================
    # 5.2 参考地图与数字正射影像图说明
    # ============================================================
    doc.add_heading('5.2 参考地图与数字正射影像图说明', level=1)

    add_para(doc,
        '视觉定位链路依赖两类参考数据：正射影像图（DOM）提供二维纹理锚点用于图像检索和匹配，'
        '数字表面模型（DSM）提供高程信息用于将二维匹配点升维为三维约束，支撑PnP位姿解算。'
    )

    add_table(doc,
        ['参考数据类型', '主要内容', '链路作用', '分辨率', '注意事项'],
        [
            ['航空正射影像', 'DJI无人机+SfM构建的2D正射影像', '提供二维纹理锚点', '0.061m/pixel', '与UAV图像时相一致，纹理质量高'],
            ['航空DSM', 'SfM技术生成的数字表面模型', '将二维匹配点升维为三维约束', '0.937m/pixel', '高程精度高，适合精确位姿解算'],
            ['卫星正射影像', 'Google Earth历史图像', '航空图不可用时的替代参考图', '0.260m/pixel', '存在时相差异，纹理质量较低'],
            ['卫星DSM', 'ALOS 30m DSM高程数据', '卫星图条件下的三维高程约束', '30m/pixel', '高程精度低，严重影响PnP求解质量'],
        ],
        caption='表5-3 参考地图与高程数据说明'
    )

    add_para(doc,
        '航空图与卫星图的分辨率、时相和高程精度差异是后续定位效果分析的关键影响因素，'
        '具体对比结果将在5.6节展开。'
    )

    # ============================================================
    # 5.3 实验平台与参数设置
    # ============================================================
    doc.add_heading('5.3 实验平台与参数设置', level=1)

    add_para(doc,
        '本实验在NVIDIA GeForce RTX 3080 (10GB) GPU平台上运行，'
        '采用CAMP作为检索模型、RoMa作为匹配模型、P3P+RANSAC作为位姿求解器。'
    )

    add_table(doc,
        ['项目', '配置'],
        [
            ['GPU', 'NVIDIA GeForce RTX 3080 (10GB)'],
            ['CPU', 'Intel Core i9-10920X @ 3.50GHz'],
            ['RAM', '128GB'],
            ['Python', '3.9.13'],
            ['PyTorch', '2.2.1+cu121'],
            ['CUDA', '12.1'],
            ['检索模型', 'CAMP (ConvNeXt-Base, University-1652预训练)'],
            ['匹配模型', 'RoMa (DINOv2-ViT-L)'],
            ['PnP求解器', 'P3P + RANSAC (OpenCV)'],
        ],
        caption='表5-4 实验平台与运行环境'
    )

    add_table(doc,
        ['参数', '值', '影响说明'],
        [
            ['TEST_INTERVAL', '1', '测试采样间隔，1=全部图像参与测试'],
            ['RETRIEVAL_COVER', '50', '检索块重叠率(%)，影响候选召回率'],
            ['RETRIEVAL_TOPN', '5', '检索Top-N数量，影响重排序策略效果'],
            ['BATCH_SIZE', '128', '特征提取批大小，影响检索速度'],
            ['resize_ratio', '0.2', '图像缩放比例，降低显存开销但可能损失精度'],
            ['coarse_res', '280', 'RoMa匹配粗分辨率（原560），降低显存从~4GB到~0.4GB'],
            ['upsample_res', '512', 'RoMa匹配上采样分辨率（原864），影响精细匹配质量'],
        ],
        caption='表5-5 定位链路关键参数设置'
    )

    add_para(doc,
        '注：RoMa匹配模型采用降低分辨率配置（coarse_res=280 vs 原始560），'
        '以适配RTX 3080的10GB显存限制。此配置可能影响精细匹配质量，在结果分析时需予以考虑。'
    )

    # ============================================================
    # 5.4 评价指标
    # ============================================================
    doc.add_heading('5.4 评价指标', level=1)

    add_para(doc,
        '本章采用以下评价指标对定位链路进行多维度评估。A@T指标反映不同阈值下的定位成功率，'
        '平均误差和中位数误差需组合使用以避免异常值影响。'
    )

    add_table(doc,
        ['指标名称', '指标含义', '评价方向', '适用实验', '对应链路环节'],
        [
            ['A@5m', '定位误差<5m的样本比例', '越高越好，反映高精度定位能力', '所有定位实验', '端到端'],
            ['A@10m', '定位误差<10m的样本比例', '越高越好', '所有定位实验', '端到端'],
            ['A@20m', '定位误差<20m的样本比例', '越高越好，反映粗定位能力', '所有定位实验', '端到端'],
            ['平均误差', '所有样本定位误差的算术平均', '越低越好，受异常值影响大', '所有定位实验', '端到端'],
            ['中位数误差', '所有样本定位误差的中位数', '越低越好，对异常值鲁棒', '所有定位实验', '端到端'],
            ['PDM@K', '检索结果与真值的空间重叠率加权评分', '越高越好，反映检索质量', '检索质量评价', '图像检索'],
            ['重投影误差', 'PnP求解后3D点在图像平面的投影偏差', '越低越好，反映位姿解算质量', 'PnP解算质量', 'PnP求解'],
            ['内点数量', 'PnP RANSAC内点数', '越多越好，反映匹配质量', '策略对比实验', 'PnP求解'],
            ['推理时间', '单帧端到端处理时间', '越低越好，影响实时性', '性能评估', '端到端'],
        ],
        caption='表5-6 评价指标说明'
    )

    add_para(doc,
        '定位误差计算公式为 e_l = sqrt((x_p - x_g)^2 + (y_p - y_g)^2)，'
        '其中(x_p, y_p)为预测坐标，(x_g, y_g)为真值坐标（UTM）。'
        'PDM@K仅用于候选检索质量评价，不应与最终米级定位误差混为一类；'
        '重投影误差和内点数量用于评估PnP位姿解算质量；'
        '推理时间用于评估链路的实时性。'
    )

    # ============================================================
    # 5.5 视觉定位算法结果分析
    # ============================================================
    doc.add_heading('5.5 视觉定位算法结果分析', level=1)

    add_para(doc,
        '基于CAMP+RoMa+Top N Re-rank的最佳组合，在航空图和卫星图上分别进行了完整测试。'
    )

    add_table(doc,
        ['参考图类型', '样本数', 'A@5m', 'A@10m', 'A@20m', '平均误差', '中位数误差'],
        [
            ['航空图', '487', '82.3%', '90.3%', '95.9%', '28.1m', '2.4m'],
            ['卫星图', '487', '13.6%', '39.4%', '63.7%', '100.0m', '13.4m'],
        ],
        caption='表5-7 视觉定位算法主结果'
    )

    add_para(doc,
        '航空图条件下定位精度显著高于卫星图（A@5m 82.3% vs 13.6%），'
        '中位数误差仅2.4m，说明大部分图像定位精度很高，少数异常值拉高了平均误差。'
    )

    add_figure(doc, FIG / 'fig5_3_pipeline.png', '图5-3 视觉定位算法完整流程可视化')

    add_para(doc,
        '图5-3展示了三类典型样本的定位流程：成功定位（误差0.09m）的链路各环节均稳定收敛；'
        '一般定位（误差9.17m）在匹配环节出现少量误匹配；定位失败（误差4797m）在检索和匹配环节均出现严重偏差。'
    )

    # ============================================================
    # 5.6 不同参考地图条件下的定位效果分析
    # ============================================================
    doc.add_heading('5.6 不同参考地图条件下的定位效果分析', level=1)

    add_para(doc,
        '航空图和卫星图两类参考地图条件造成定位性能差异显著，'
        '主要原因包括影像分辨率、DSM分辨率、时相差异和纹理质量。'
    )

    add_table(doc,
        ['参考图类型', '影像分辨率', 'DSM分辨率', 'A@5m', 'A@10m', 'A@20m', '平均误差'],
        [
            ['航空图', '0.061m/pixel', '0.937m', '82.3%', '90.3%', '95.9%', '28.1m'],
            ['卫星图', '0.260m/pixel', '30m', '13.6%', '39.4%', '63.7%', '100.0m'],
        ],
        caption='表5-8 不同参考地图条件下的定位效果'
    )

    add_figure(doc, FIG / 'fig5_4_high_vs_low.png', '图5-4 航空图与卫星图定位效果对比')

    add_para(doc,
        '卫星图条件下A@5m仅为13.6%，但A@20m达到63.7%，说明链路在卫星图条件下仍具备粗定位能力。'
        '低精度并非算法失效，而是参考图分辨率（0.260m vs 0.061m）、DSM分辨率（30m vs 0.937m）'
        '和时相差异共同导致的链路性能下降。'
    )

    # ============================================================
    # 5.7 不同飞行高度和视角条件下的定位效果分析
    # ============================================================
    doc.add_heading('5.7 不同飞行高度和视角条件下的定位效果分析', level=1)

    add_para(doc,
        '飞行俯仰角和高度对定位精度有显著影响。低空多视角条件下，'
        '几何视角差异直接影响与正射参考图的相似性。'
    )

    add_table(doc,
        ['分析维度', '分组范围', '样本数', 'A@5m', 'A@20m', '平均误差', '解释重点'],
        [
            ['俯仰角', '20°-50°（倾斜）', '301', '79.4%', '96.7%', '35.5m', '倾斜视角透视畸变严重'],
            ['俯仰角', '50°-70°（过渡）', '22', '63.6%', '90.9%', '17.8m', '样本量有限，结论仅供参考'],
            ['俯仰角', '70°-90°（正下视）', '164', '90.2%', '95.1%', '16.1m', '正下视精度最高'],
            ['高度', '30-100m', '48', '33.3%', '87.5%', '185.7m', '低高度样本误差大'],
            ['高度', '100-200m', '439', '87.7%', '96.8%', '10.9m', '主要飞行高度区间'],
            ['高度', '200-300m', '0', '—', '—', '—', 'Demo子集无此高度样本'],
            ['场景', 'QZ_SongCity', '286', '98.3%', '100.0%', '1.73m', '纹理清晰，定位稳定'],
            ['场景', 'Qingzhou_3_2', '59', '96.6%', '100.0%', '3.04m', '纹理清晰，定位稳定'],
            ['场景', 'QingZhou_2024', '142', '44.4%', '85.9%', '91.75m', '存在大面积水域，弱纹理'],
        ],
        caption='表5-9 不同高度与视角条件下定位效果'
    )

    add_figure(doc, FIG / 'fig5_5_height_pitch_analysis.png', '图5-5 不同高度与视角条件下定位效果分析')

    add_para(doc,
        '正下视（70°-90°）A@5m达到90.2%，倾斜视角（20°-50°）降至79.4%，'
        '证实了低空多视角条件的核心挑战。50°-70°分组样本量仅22张，统计意义有限。'
        'QingZhou_2024场景表现异常（A@5m=44.4%），主要因存在大面积水域导致弱纹理匹配困难。'
    )

    # ============================================================
    # 5.8 定位策略与匹配方法对比分析
    # ============================================================
    doc.add_heading('5.8 定位策略与匹配方法对比分析', level=1)

    add_para(doc,
        '本节比较不同定位策略、匹配方法对定位稳定性的影响，筛选当前链路中最关键的性能影响因素。'
    )

    add_table(doc,
        ['分析对象', '对比设置', 'A@5m', 'A@10m', 'A@20m', '主要结论'],
        [
            ['定位策略', 'Top N Re-rank', '82.3%', '90.3%', '95.9%', '精度最高，综合最优'],
            ['定位策略', 'Top1', '68.8%', '78.2%', '81.9%', '最简单但容错率低'],
            ['定位策略', 'Most Inliers', '79.5%', '86.2%', '87.3%', '受异常值影响较大'],
            ['匹配方法', 'RoMa（密集/学习）', '82.3%', '90.3%', '95.9%', '密集匹配精度最高'],
            ['匹配方法', 'SIFT（稀疏/手工）', '31.2%', '35.9%', '39.2%', '稀疏手工方法差距巨大'],
            ['先验噪声', 'Yaw std=0°', '74.6%', '87.6%', '94.2%', '无噪声基准（论文数据）'],
            ['先验噪声', 'Yaw std=30°', '70.5%', '83.5%', '90.2%', '中等噪声影响'],
            ['先验噪声', 'Yaw std=60°', '48.9%', '63.0%', '70.0%', '严重噪声导致精度崩溃'],
        ],
        caption='表5-10 定位策略、匹配方法与先验噪声对比'
    )

    add_para(doc,
        '注：先验噪声数据来源于AnyVisLoc论文Table 7，非本文复现实验结果。'
    )

    add_figure(doc, FIG / 'fig5_6_strategy_matching.png', '图5-6 定位策略、匹配方法与先验噪声对比分析')

    add_para(doc,
        'Top N Re-rank策略（A@5m=82.3%）显著优于Top1策略（68.8%），'
        '通过综合Top-5检索结果的匹配内点数进行重排序，有效提高了定位鲁棒性。'
        'RoMa密集匹配（A@5m=82.3%）大幅优于SIFT稀疏手工方法（31.2%），差距达51.1个百分点。'
        'Yaw噪声标准差超过10°时精度开始显著下降，超过60°时A@5m降至48.9%。'
    )

    # ============================================================
    # 5.9 与目标检测结果联动的定位误差分析框架
    # ============================================================
    doc.add_heading('5.9 与目标检测结果联动的定位误差分析框架', level=1)

    add_para(doc,
        '本节提出目标检测框与定位链路联动的误差分析框架，说明检测误差如何传播至地理坐标。'
        '需要说明的是，本节是分析框架而非已完成量化实验。'
    )

    add_table(doc,
        ['检测框输入类型', '扰动设置', '所需参数', '输出指标', '需回答的问题'],
        [
            ['人工标注框', '无扰动（理想情况）', '目标像素坐标真值', '基准定位误差', '理想检测下的定位精度上限'],
            ['模型预测框', '检测网络输出', '置信度、类别、框坐标', '定位误差、置信度相关性', '检测精度如何影响定位精度'],
            ['扰动框', '像素偏移±Δpx', '偏移量、图像尺寸', '地理误差随像素偏移的变化', '框中心偏移产生多大地理误差'],
            ['多目标框', '多个检测结果', '目标间距、类别', '多目标定位误差分布', '多目标场景下的定位鲁棒性'],
        ],
        caption='表5-11 检测框联动定位误差分析框架（分析框架模板，非实验结果）'
    )

    add_figure(doc, FIG / 'fig5_7_error_propagation.png', '图5-7 检测框→定位误差传播链路')

    add_para(doc,
        '误差传播路径为：检测框中心偏移→像素误差→射线角度误差→地面投影误差→目标坐标误差。'
        '在缺少目标级真值的条件下，本节仅完成误差传播链路分析框架设计，'
        '后续需结合红外/可见光检测结果进行量化验证。'
    )

    # ============================================================
    # 5.10 典型定位结果与失败案例分析
    # ============================================================
    doc.add_heading('5.10 典型定位结果与失败案例分析', level=1)

    add_para(doc,
        '通过成功案例和失败案例解释视觉定位链路何时稳定、何时失效，'
        '以及失败主要来自视角退化、参考图差异还是检测框联动风险。'
    )

    add_table(doc,
        ['案例类型', '场景特征', '结果现象', '主要原因', '正文分析重点'],
        [
            ['成功案例', '正下视+航空图+清晰纹理', '误差<2m，内点>1000', '视角匹配度高，DSM约束强', '链路在理想条件下可达米级精度'],
            ['视角退化失败', '大倾斜角(20°-30°)', '误差>100m，内点少', '透视畸变，与正射参考图相似度低', '低空多视角条件的核心挑战'],
            ['参考图差异失败', '卫星图+低DSM分辨率', '误差>50m', 'DSM分辨率30m导致PnP约束弱', '参考图质量对链路的决定性影响'],
            ['检测框联动风险', '检测框偏移±Δpx', '地理误差线性放大', '框中心误差→射线角度误差', '不能与UAV定位链路结果混为一谈'],
        ],
        caption='表5-12 典型案例选择与失败原因说明'
    )

    add_figure(doc, FIG / 'fig5_8_success_failure.png', '图5-8 典型定位结果与失败案例分析')

    add_para(doc,
        '成功案例中，QZ_SongCity场景286张图像的A@5m达到98.3%，平均误差仅1.73m，'
        '证实链路在理想条件下可达米级精度。失败案例主要集中在QingZhou_2024场景，'
        '该场景存在大面积水域和大倾斜角拍摄，导致检索和匹配均出现严重偏差。'
    )

    # ============================================================
    # 5.11 本章小结
    # ============================================================
    doc.add_heading('5.11 本章小结', level=1)

    add_para(doc,
        '本章基于AnyVisLoc数据集Demo子集，对CAMP+RoMa+PnP构成的低空多视角UAV视觉定位链路'
        '进行了系统性测试。主要结论如下：'
    )

    doc.add_paragraph(
        '数据与场景层面：本章使用AnyVisLoc原方法配套定位数据和Demo测试子集，'
        '主要服务于UAV视觉定位链路测试。Demo子集飞行高度33-193m，俯仰角20°-90°，'
        '覆盖正下视到倾斜视角的多种条件。',
        style='List Bullet'
    )

    doc.add_paragraph(
        '定位链路层面：CAMP+RoMa+PnP构成当前基础链路，结果应围绕检索候选、'
        '匹配质量和PnP内点稳定性展开。航空图条件下A@5m=82.3%，中位数误差2.4m，'
        '链路在理想条件下可达米级精度。',
        style='List Bullet'
    )

    doc.add_paragraph(
        '参考地图层面：航空图条件（A@5m=82.3%）显著优于卫星图条件（A@5m=13.6%），'
        '主要受限于卫星图DSM分辨率（30m vs 0.937m）和时相差异。'
        '参考图质量是影响定位精度的首要因素。',
        style='List Bullet'
    )

    doc.add_paragraph(
        '飞行视角与高度层面：正下视（70°-90°）A@5m=90.2%，倾斜视角（20°-50°）A@5m=79.4%，'
        '倾斜拍摄带来的透视畸变是低空多视角定位的核心挑战。'
        '50°-70°分组样本量仅22张，统计意义有限。',
        style='List Bullet'
    )

    doc.add_paragraph(
        '策略匹配层面：Top N Re-rank策略（A@5m=82.3%）在精度、速度、内存间取得最佳平衡，'
        'RoMa密集匹配（A@5m=82.3%）大幅优于SIFT稀疏手工方法（31.2%）。'
        '先验噪声控制（Yaw<10°）是后续优化重点。',
        style='List Bullet'
    )

    doc.add_paragraph(
        '任务边界层面：检测框联动定位目前属于误差传播分析框架，'
        '不能写成已经完成的水面人体目标定位实测。后续需结合红外/可见光检测结果'
        '进行量化验证，并引入EKF时序平滑消除单帧噪声。',
        style='List Bullet'
    )

    # 保存
    output_path = BASE / 'reports' / '第五章_实验验证与结果分析.docx'
    doc.save(str(output_path))
    print(f'第五章已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_chapter5()
