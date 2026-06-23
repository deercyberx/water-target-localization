"""检查poi-tl生成的docx格式规范"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

DOC_PATH = r"C:\Users\deer\Desktop\Projects\Research\projects\水上目标定位\reports\第五章_实验验证与结果分析_poi-tl_final.docx"

doc = Document(DOC_PATH)

issues = []
passed = []

# ============================================================
# 1. 标题层级检查
# ============================================================
print("=" * 60)
print("1. 标题层级检查")
print("=" * 60)

heading_counts = {}
heading_samples = {}
for para in doc.paragraphs:
    style_name = para.style.name if para.style else "None"
    if style_name.startswith("Heading"):
        heading_counts[style_name] = heading_counts.get(style_name, 0) + 1
        if style_name not in heading_samples:
            heading_samples[style_name] = []
        if len(heading_samples[style_name]) < 3:
            heading_samples[style_name].append(para.text[:60])

print(f"  标题分布: {heading_counts}")
for style, samples in heading_samples.items():
    for s in samples:
        print(f"    [{style}] {s}")

# 检查是否有Heading 1用于节标题
h1_count = heading_counts.get("Heading 1", 0)
h2_count = heading_counts.get("Heading 2", 0)
h3_count = heading_counts.get("Heading 3", 0)

if h1_count > 0:
    passed.append(f"有Heading 1节标题 ({h1_count}个)")
else:
    issues.append("缺少Heading 1节标题")

if h3_count > 0:
    passed.append(f"有Heading 3图表标题 ({h3_count}个)")
else:
    issues.append("缺少Heading 3图表标题")

# ============================================================
# 2. 字体检查
# ============================================================
print("\n" + "=" * 60)
print("2. 字体检查")
print("=" * 60)

font_issues = []
body_fonts = set()
heading_fonts = set()
body_sizes = set()
heading_sizes = set()

for para in doc.paragraphs:
    style_name = para.style.name if para.style else "None"
    is_heading = style_name.startswith("Heading")

    for run in para.runs:
        if not run.text.strip():
            continue

        # 检查字体
        font_name = run.font.name
        # 也检查东亚字体
        rPr = run._element.find(qn('w:rPr'))
        ea_font = None
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea_font = rFonts.get(qn('w:eastAsia'))

        effective_font = ea_font or font_name

        if is_heading:
            if effective_font:
                heading_fonts.add(effective_font)
        else:
            if effective_font:
                body_fonts.add(effective_font)

        # 检查字号
        size = run.font.size
        if size:
            pt_val = size.pt
            if is_heading:
                heading_sizes.add(pt_val)
            else:
                body_sizes.add(pt_val)

print(f"  正文字体集合: {body_fonts}")
print(f"  正文字号集合: {body_sizes}")
print(f"  标题字体集合: {heading_fonts}")
print(f"  标题字号集合: {heading_sizes}")

# 检查正文是否为宋体12pt
if body_fonts:
    has_songti = any("宋体" in f or "SimSun" in f or "simsun" in f for f in body_fonts)
    if has_songti:
        passed.append(f"正文字体包含宋体")
    else:
        issues.append(f"正文字体不是宋体，实际: {body_fonts}")

if body_sizes:
    if 12.0 in body_sizes:
        passed.append("正文字号包含12pt")
    else:
        issues.append(f"正文字号不是12pt，实际: {body_sizes}")

# 检查标题是否为黑体
if heading_fonts:
    has_heiti = any("黑体" in f or "SimHei" in f or "simhei" in f for f in heading_fonts)
    if has_heiti:
        passed.append("标题字体包含黑体")
    else:
        issues.append(f"标题字体不是黑体，实际: {heading_fonts}")
else:
    issues.append("未检测到标题字体信息")

# ============================================================
# 3. 行距检查
# ============================================================
print("\n" + "=" * 60)
print("3. 行距检查")
print("=" * 60)

line_spacing_values = set()
line_spacing_rules = set()
sample_count = 0

for para in doc.paragraphs:
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        # line_spacing 可以是 float (倍数) 或 Length (固定值)
        from docx.shared import Length
        if isinstance(pf.line_spacing, (int, float)):
            line_spacing_values.add(float(pf.line_spacing))
        else:
            # It's a Length object
            line_spacing_values.add(f"{pf.line_spacing.pt}pt (固定值)")

    if pf.line_spacing_rule is not None:
        line_spacing_rules.add(str(pf.line_spacing_rule))

print(f"  行距值: {line_spacing_values}")
print(f"  行距规则: {line_spacing_rules}")

# 检查是否为1.5倍行距
has_1_5 = False
for v in line_spacing_values:
    if isinstance(v, float) and abs(v - 1.5) < 0.01:
        has_1_5 = True
        break

if has_1_5:
    passed.append("行距包含1.5倍")
else:
    issues.append(f"行距不是1.5倍，实际值: {line_spacing_values}")

# ============================================================
# 4. 页边距检查
# ============================================================
print("\n" + "=" * 60)
print("4. 页边距检查")
print("=" * 60)

for section in doc.sections:
    top = section.top_margin
    bottom = section.bottom_margin
    left = section.left_margin
    right = section.right_margin

    top_cm = top / 914400 * 2.54 * 10 if top else None  # EMU to cm
    bottom_cm = bottom / 914400 * 2.54 * 10 if bottom else None
    left_cm = left / 914400 * 2.54 * 10 if left else None
    right_cm = right / 914400 * 2.54 * 10 if right else None

    # 更准确的EMU到cm转换: 1 cm = 360000 EMU
    if top:
        top_cm = top / 360000
    if bottom:
        bottom_cm = bottom / 360000
    if left:
        left_cm = left / 360000
    if right:
        right_cm = right / 360000

    print(f"  上: {top_cm:.2f}cm, 下: {bottom_cm:.2f}cm, 左: {left_cm:.2f}cm, 右: {right_cm:.2f}cm")

    def check_margin(name, actual, expected, tolerance=0.05):
        if actual is None:
            issues.append(f"页边距{name}未设置")
            return
        if abs(actual - expected) > tolerance:
            issues.append(f"页边距{name}不符: 期望{expected}cm, 实际{actual:.2f}cm")
        else:
            passed.append(f"页边距{name}符合 ({actual:.2f}cm)")

    check_margin("上", top_cm, 2.54)
    check_margin("下", bottom_cm, 2.54)
    check_margin("左", left_cm, 3.17)
    check_margin("右", right_cm, 3.17)

# ============================================================
# 5. 表格格式检查
# ============================================================
print("\n" + "=" * 60)
print("5. 表格格式检查")
print("=" * 60)

table_count = len(doc.tables)
print(f"  表格总数: {table_count}")

for i, table in enumerate(doc.tables):
    print(f"\n  表格 {i+1}: {len(table.rows)}行 x {len(table.columns)}列")

    # 检查表头行
    if len(table.rows) > 0:
        header_row = table.rows[0]
        header_bg_colors = []
        for cell in header_row.cells:
            # 检查背景色
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    color = shd.get(qn('w:color'))
                    val = shd.get(qn('w:val'))
                    header_bg_colors.append(f"fill={fill},color={color},val={val}")

        if header_bg_colors:
            print(f"    表头背景: {header_bg_colors[:3]}...")
            if any(c and "fill=None" not in c for c in header_bg_colors):
                passed.append(f"表格{i+1}表头有背景色")
            else:
                issues.append(f"表格{i+1}表头无背景色")
        else:
            issues.append(f"表格{i+1}表头无背景色信息")

        # 检查数据居中
        center_count = 0
        total_data_cells = 0
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        total_data_cells += 1
                        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            center_count += 1

        if total_data_cells > 0:
            center_ratio = center_count / total_data_cells
            print(f"    居中段落: {center_count}/{total_data_cells} ({center_ratio:.1%})")
            if center_ratio > 0.5:
                passed.append(f"表格{i+1}数据大多居中 ({center_ratio:.0%})")
            else:
                issues.append(f"表格{i+1}数据居中比例低 ({center_ratio:.0%})")

# ============================================================
# 6. 图片格式检查
# ============================================================
print("\n" + "=" * 60)
print("6. 图片格式检查")
print("=" * 60)

# 通过检查inline shapes和段落对齐来判断图片
inline_shapes = doc.inline_shapes
print(f"  内联图片数: {len(inline_shapes)}")

# 检查包含图片的段落是否居中
img_para_count = 0
img_center_count = 0
for para in doc.paragraphs:
    # 检查段落是否包含图片
    has_image = False
    for run in para.runs:
        if run._element.find(qn('w:drawing')) is not None:
            has_image = True
            break
        # 也检查旧式图片
        if run._element.find(qn('w:pict')) is not None:
            has_image = True
            break

    if has_image:
        img_para_count += 1
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            img_center_count += 1
        # 检查图片宽度
        for run in para.runs:
            drawing = run._element.find(qn('w:drawing'))
            if drawing is not None:
                # 尝试获取图片宽度
                inline = drawing.find(qn('wp:inline'))
                if inline is not None:
                    extent = inline.find(qn('wp:extent'))
                    if extent is not None:
                        cx = extent.get('cx')
                        cy = extent.get('cy')
                        if cx:
                            width_cm = int(cx) / 360000
                            print(f"    图片宽度: {width_cm:.2f}cm")

print(f"  含图片段落: {img_para_count}, 居中: {img_center_count}")

if img_para_count > 0:
    if img_center_count == img_para_count:
        passed.append(f"所有图片段落居中 ({img_center_count}/{img_para_count})")
    else:
        issues.append(f"部分图片未居中: {img_center_count}/{img_para_count}")
else:
    print("  (未检测到图片段落，可能使用其他方式嵌入)")

# ============================================================
# 7. 段落间距检查
# ============================================================
print("\n" + "=" * 60)
print("7. 段落间距检查")
print("=" * 60)

spacing_before = set()
spacing_after = set()

for para in doc.paragraphs:
    pf = para.paragraph_format
    if pf.space_before is not None:
        spacing_before.add(pf.space_before.pt)
    if pf.space_after is not None:
        spacing_after.add(pf.space_after.pt)

print(f"  段前间距值: {sorted(spacing_before)}")
print(f"  段后间距值: {sorted(spacing_after)}")

# 一般要求段前段后为0或适当值
if spacing_before:
    passed.append(f"段前间距已设置: {sorted(spacing_before)}")
else:
    issues.append("段前间距未设置")

if spacing_after:
    passed.append(f"段后间距已设置: {sorted(spacing_after)}")
else:
    issues.append("段后间距未设置")

# ============================================================
# 总结
# ============================================================
print("\n" + "=" * 60)
print("总结")
print("=" * 60)

print(f"\n通过项 ({len(passed)}):")
for p in passed:
    print(f"  [PASS] {p}")

print(f"\n问题项 ({len(issues)}):")
for issue in issues:
    print(f"  [FAIL] {issue}")

if not issues:
    print("\n结论: 全部通过")
else:
    print(f"\n结论: 不通过 (共{len(issues)}个问题)")
