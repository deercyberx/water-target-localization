# -*- coding: utf-8 -*-
"""
修复技术报告格式问题（素材清单 F1-F8）
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

INPUT_FILE = '水面目标精准定位技术报告_生成版_v3.docx'
OUTPUT_FILE = '水面目标精准定位技术报告_终稿.docx'

def fix_paper_size(doc):
    """F7: 纸张大小从 US Letter 改为 A4"""
    print('[F7] 修复纸张大小: US Letter → A4 (210×297mm)')
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # 调整边距
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

def fix_heading1_style(doc):
    """F4: Heading 1 样式 - 粗体/居中/宋体/行距2.408倍"""
    print('[F4] 修复 Heading 1 样式')
    style = doc.styles['Heading 1']
    font = style.font
    font.bold = True
    font.size = Pt(16)
    font.name = 'Times New Roman'
    # 设置中文字体
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn('w:eastAsia'), '宋体')

    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.line_spacing = 2.408

def fix_heading2_style(doc):
    """F5: Heading 2 样式 - 12pt/宋体"""
    print('[F5] 修复 Heading 2 样式')
    style = doc.styles['Heading 2']
    font = style.font
    font.bold = True
    font.size = Pt(12)
    font.name = 'Times New Roman'
    # 设置中文字体
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn('w:eastAsia'), '宋体')

    pf = style.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

def fix_heading3_style(doc):
    """修复 Heading 3 样式"""
    print('[补充] 修复 Heading 3 样式')
    style = doc.styles['Heading 3']
    font = style.font
    font.bold = True
    font.size = Pt(12)
    font.name = 'Times New Roman'
    rpr = style.element.find(qn('w:rPr'))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn('w:eastAsia'), '宋体')

    pf = style.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5

def fix_normal_style(doc):
    """F6: 正文行距改为 1.5 倍"""
    print('[F6] 修复正文行距: 1.5倍行距')
    style = doc.styles['Normal']
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

def fix_heading_paragraphs(doc):
    """修复所有 Heading 段落的实际格式"""
    print('[F3/F4] 修复 Heading 段落格式')
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.408
            # 确保字体
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(16)
        elif p.style.name == 'Heading 2':
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        elif p.style.name == 'Heading 3':
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(12)

def fix_normal_paragraphs(doc):
    """修复正文段落格式"""
    print('[F6] 修复正文段落格式')
    for p in doc.paragraphs:
        if p.style.name == 'Normal' and p.text.strip():
            p.paragraph_format.line_spacing = 1.5

def fix_chapter3_heading(doc):
    """F3: 检查第三章标题级别 - 只修正明确的章节标题"""
    print('[F3] 检查第三章标题级别')
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # 只匹配以"第三章"开头的标题段落，不匹配包含"第三章"的正文
        if text.startswith('第三章') and len(text) < 50:
            print(f'  段落 {i}: [{p.style.name}] {text[:60]}')
            if p.style.name != 'Heading 1':
                print(f'  → 修正为 Heading 1')
                p.style = doc.styles['Heading 1']

def add_toc(doc):
    """F1: 添加目录"""
    print('[F1] 添加目录页')
    # 找到"目  录"段落
    for i, p in enumerate(doc.paragraphs):
        if '目' in p.text and '录' in p.text and len(p.text.strip()) < 10:
            # 在此段落后添加 TOC 域代码
            # 清除原有内容
            for run in p.runs:
                run.clear()
            p.text = ''

            # 添加 TOC 域代码
            fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run1 = p.add_run()
            run1._r.append(fldChar_begin)

            instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
            run2 = p.add_run()
            run2._r.append(instrText)

            fldChar_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
            run3 = p.add_run()
            run3._r.append(fldChar_separate)

            # 添加占位文本
            run4 = p.add_run('【请右键更新域以生成目录】')
            run4.font.color.rgb = None

            fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run5 = p.add_run()
            run5._r.append(fldChar_end)

            print(f'  在段落 {i} 添加 TOC 域代码')
            break

def main():
    print(f'读取: {INPUT_FILE}')
    doc = Document(INPUT_FILE)

    print('=' * 50)
    print('开始修复格式问题')
    print('=' * 50)

    # F7: 纸张大小
    fix_paper_size(doc)

    # F4: Heading 1 样式
    fix_heading1_style(doc)

    # F5: Heading 2 样式
    fix_heading2_style(doc)

    # Heading 3 样式
    fix_heading3_style(doc)

    # F6: 正文行距
    fix_normal_style(doc)

    # F3: 第三章标题级别
    fix_chapter3_heading(doc)

    # 修复段落实际格式
    fix_heading_paragraphs(doc)
    fix_normal_paragraphs(doc)

    # F1: 添加目录
    add_toc(doc)

    print('=' * 50)
    print(f'保存: {OUTPUT_FILE}')
    doc.save(OUTPUT_FILE)
    print('完成!')

if __name__ == '__main__':
    main()
