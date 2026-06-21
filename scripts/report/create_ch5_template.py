# -*- coding: utf-8 -*-
"""
创建第五章 poi-tl 模板
用 python-docx 生成带 {{变量}} 占位符的模板文件
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = BASE / 'poi-tl-project' / 'poi-tl' / 'templates'
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

def set_heading_style(doc):
    """设置标题样式"""
    for level, size in [(1, 16), (2, 14), (3, 12)]:
        style = doc.styles[f'Heading {level}']
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.name = 'SimHei'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

def add_table_placeholder(doc, caption, headers, rows_sample):
    """添加表格占位符"""
    doc.add_paragraph(f'{{{{{caption}}}}}', style='Normal')
    # 示例表格结构
    table = doc.add_table(rows=1 + len(rows_sample), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows_sample):
        for c, val in enumerate(row):
            table.rows[r+1].cells[c].text = str(val)

def create_template():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = 'SimSun'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    set_heading_style(doc)

    # ============================================================
    # 第五章标题
    # ============================================================
    doc.add_heading('第五章 实验验证与结果分析', level=1)

    # ============================================================
    # 5.1 实验数据与测试场景
    # ============================================================
    doc.add_heading('5.1 实验数据与测试场景', level=2)

    doc.add_paragraph('{{5_1_intro}}')

    doc.add_paragraph('表5-1 AnyVisLoc数据集属性', style='Heading 3')
    add_table_placeholder(doc, 'table_5_1',
        ['属性', '内容'],
        [['数据集名称', '{{ds_name}}'], ['UAV图像数量', '{{ds_images}}']]
    )

    doc.add_paragraph('表5-2 Demo测试子集说明', style='Heading 3')
    add_table_placeholder(doc, 'table_5_2',
        ['测试区域', '查询图像数', '航空参考图', '卫星参考图'],
        [['{{scene_name}}', '{{scene_count}}', '{{scene_aerial}}', '{{scene_satellite}}']]
    )

    doc.add_paragraph('{{5_1_figure_ref}}')
    doc.add_paragraph('{{@fig5_1}}')

    doc.add_paragraph('{{5_1_conclusion}}')

    # ============================================================
    # 5.2 参考地图与数字正射影像图说明
    # ============================================================
    doc.add_heading('5.2 参考地图与数字正射影像图说明', level=2)

    doc.add_paragraph('{{5_2_intro}}')

    doc.add_paragraph('表5-3 参考地图与高程数据说明', style='Heading 3')
    add_table_placeholder(doc, 'table_5_3',
        ['参考数据类型', '主要内容', '链路作用', '分辨率', '注意事项'],
        [['{{ref_type}}', '{{ref_content}}', '{{ref_role}}', '{{ref_res}}', '{{ref_note}}']]
    )

    doc.add_paragraph('{{5_2_conclusion}}')

    # ============================================================
    # 5.3 实验平台与参数设置
    # ============================================================
    doc.add_heading('5.3 实验平台与参数设置', level=2)

    doc.add_paragraph('{{5_3_intro}}')

    doc.add_paragraph('表5-4 实验平台与运行环境说明', style='Heading 3')
    add_table_placeholder(doc, 'table_5_4',
        ['项目', '配置'],
        [['{{env_item}}', '{{env_value}}']]
    )

    doc.add_paragraph('表5-5 定位链路关键参数设置', style='Heading 3')
    add_table_placeholder(doc, 'table_5_5',
        ['参数', '值', '影响说明'],
        [['{{param_name}}', '{{param_value}}', '{{param_effect}}']]
    )

    doc.add_paragraph('{{5_3_note}}')

    # ============================================================
    # 5.4 评价指标
    # ============================================================
    doc.add_heading('5.4 评价指标', level=2)

    doc.add_paragraph('{{5_4_intro}}')

    doc.add_paragraph('表5-6 评价指标说明', style='Heading 3')
    add_table_placeholder(doc, 'table_5_6',
        ['指标名称', '指标含义', '评价方向', '适用实验'],
        [['{{metric_name}}', '{{metric_meaning}}', '{{metric_direction}}', '{{metric_usage}}']]
    )

    doc.add_paragraph('{{5_4_conclusion}}')

    # ============================================================
    # 5.5 视觉定位算法结果分析
    # ============================================================
    doc.add_heading('5.5 视觉定位算法结果分析', level=2)

    doc.add_paragraph('{{5_5_intro}}')

    doc.add_paragraph('表5-7 视觉定位算法主结果', style='Heading 3')
    add_table_placeholder(doc, 'table_5_7',
        ['参考图类型', '样本数', 'A@5m', 'A@10m', 'A@20m', '平均误差', '中位数误差'],
        [['{{result_type}}', '{{result_n}}', '{{result_a5}}', '{{result_a10}}', '{{result_a20}}', '{{result_mean}}', '{{result_median}}']]
    )

    doc.add_paragraph('{{5_5_analysis}}')

    doc.add_paragraph('{{@fig5_3}}')

    doc.add_paragraph('{{5_5_conclusion}}')

    # ============================================================
    # 5.6 不同参考地图条件下的定位效果分析
    # ============================================================
    doc.add_heading('5.6 不同参考地图条件下的定位效果分析', level=2)

    doc.add_paragraph('{{5_6_intro}}')

    doc.add_paragraph('表5-8 不同参考地图条件下的定位效果', style='Heading 3')
    add_table_placeholder(doc, 'table_5_8',
        ['参考图类型', '影像分辨率', 'DSM分辨率', 'A@5m', 'A@10m', 'A@20m', '平均误差'],
        [['{{map_type}}', '{{map_img_res}}', '{{map_dsm_res}}', '{{map_a5}}', '{{map_a10}}', '{{map_a20}}', '{{map_mean}}']]
    )

    doc.add_paragraph('{{@fig5_4}}')

    doc.add_paragraph('{{5_6_conclusion}}')

    # ============================================================
    # 5.7 不同飞行高度和视角条件下的定位效果分析
    # ============================================================
    doc.add_heading('5.7 不同飞行高度和视角条件下的定位效果分析', level=2)

    doc.add_paragraph('{{5_7_intro}}')

    doc.add_paragraph('表5-9 不同高度与视角条件下定位效果', style='Heading 3')
    add_table_placeholder(doc, 'table_5_9',
        ['分析维度', '分组范围', '样本数', 'A@5m', 'A@20m', '中位数误差', '解释重点'],
        [['{{dim_name}}', '{{dim_range}}', '{{dim_n}}', '{{dim_a5}}', '{{dim_a20}}', '{{dim_median}}', '{{dim_note}}']]
    )

    doc.add_paragraph('{{@fig5_5}}')

    doc.add_paragraph('{{5_7_conclusion}}')

    # ============================================================
    # 5.8 定位策略与匹配方法对比分析
    # ============================================================
    doc.add_heading('5.8 定位策略与匹配方法对比分析', level=2)

    doc.add_paragraph('{{5_8_intro}}')

    doc.add_paragraph('表5-10 定位策略、匹配方法与先验噪声对比', style='Heading 3')
    add_table_placeholder(doc, 'table_5_10',
        ['分析对象', '对比设置', 'A@5m', 'A@10m', 'A@20m', '主要结论'],
        [['{{compare_obj}}', '{{compare_setting}}', '{{compare_a5}}', '{{compare_a10}}', '{{compare_a20}}', '{{compare_conclusion}}']]
    )

    doc.add_paragraph('{{@fig5_6}}')

    doc.add_paragraph('{{5_8_conclusion}}')

    # ============================================================
    # 5.9 与目标检测结果联动的定位误差分析框架
    # ============================================================
    doc.add_heading('5.9 与目标检测结果联动的定位误差分析框架', level=2)

    doc.add_paragraph('{{5_9_intro}}')

    doc.add_paragraph('表5-11 检测框联动定位误差分析框架', style='Heading 3')
    add_table_placeholder(doc, 'table_5_11',
        ['检测框输入类型', '扰动设置', '所需参数', '输出指标', '需回答的问题'],
        [['{{detect_type}}', '{{detect_perturbation}}', '{{detect_params}}', '{{detect_output}}', '{{detect_question}}']]
    )

    doc.add_paragraph('{{@fig5_7}}')

    doc.add_paragraph('{{5_9_conclusion}}')

    # ============================================================
    # 5.10 典型定位结果与失败案例分析
    # ============================================================
    doc.add_heading('5.10 典型定位结果与失败案例分析', level=2)

    doc.add_paragraph('{{5_10_intro}}')

    doc.add_paragraph('表5-12 典型案例选择与失败原因说明', style='Heading 3')
    add_table_placeholder(doc, 'table_5_12',
        ['案例类型', '场景特征', '结果现象', '主要原因', '正文分析重点'],
        [['{{case_type}}', '{{case_feature}}', '{{case_result}}', '{{case_reason}}', '{{case_analysis}}']]
    )

    doc.add_paragraph('{{@fig5_8}}')

    doc.add_paragraph('{{5_10_conclusion}}')

    # ============================================================
    # 5.11 本章小结
    # ============================================================
    doc.add_heading('5.11 本章小结', level=2)

    doc.add_paragraph('{{5_11_summary}}')

    doc.add_paragraph('{{5_11_conclusion_1}}')
    doc.add_paragraph('{{5_11_conclusion_2}}')
    doc.add_paragraph('{{5_11_conclusion_3}}')
    doc.add_paragraph('{{5_11_conclusion_4}}')

    # 保存
    output_path = TEMPLATE_DIR / 'chapter5_template.docx'
    doc.save(str(output_path))
    print(f'模板已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_template()
