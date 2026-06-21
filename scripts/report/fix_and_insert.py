# -*- coding: utf-8 -*-
"""
格式修复 + 图表插入脚本
输入: 水面目标精准定位技术报告_终稿_v3.docx
输出: 水面目标精准定位技术报告_终稿_final.docx
"""
import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

INPUT_FILE = '水面目标精准定位技术报告_终稿_v3.docx'
OUTPUT_FILE = '水面目标精准定位技术报告_终稿_final.docx'
IMG_DIR = '素材/截图'

# ──────────────────────────────────────────────
# 图表配置
# ──────────────────────────────────────────────
CHARTS = [
    {
        'file': 'error_histogram.png',
        'chapter': '5.5',
        'fig_num': '5-1',
        'caption': '定位误差分布直方图',
        'ref_text': '如图5-1所示，航空图定位误差呈现明显的长尾分布特征，大部分图像的定位误差集中在5米以内。',
    },
    {
        'file': 'error_cdf.png',
        'chapter': '5.5',
        'fig_num': '5-2',
        'caption': '定位误差累积分布函数',
        'ref_text': '图5-2展示了定位误差的累积分布函数（CDF），可直观观察不同误差阈值下的样本比例。',
    },
    {
        'file': 'strategy_comparison.png',
        'chapter': '5.8',
        'fig_num': '5-3',
        'caption': '不同定位策略精度对比',
        'ref_text': '如图5-3所示，Top-N Re-rank策略在各精度阈值下均优于Top-1和Most Inliers策略。',
    },
    {
        'file': 'pitch_impact.png',
        'chapter': '5.7',
        'fig_num': '5-4',
        'caption': '俯仰角对定位精度的影响',
        'ref_text': '如图5-4所示，正下视拍摄（70°-90°）的定位精度最高，随俯仰角减小精度逐步下降。',
    },
    {
        'file': 'scene_comparison.png',
        'chapter': '5.6',
        'fig_num': '5-5',
        'caption': '分场景定位精度对比',
        'ref_text': '如图5-5所示，航空图与卫星图在各场景下的定位精度差异显著，航空图整体精度远高于卫星图。',
    },
    {
        'file': 'qingzhou_2024_sample_1.jpg',
        'chapter': '5.10',
        'fig_num': '5-6',
        'caption': 'QingZhou_2024场景示例',
        'ref_text': '如图5-6展示了QingZhou_2024场景的典型无人机图像，该场景包含大倾斜角拍摄和弱纹理区域。',
    },
    {
        'file': 'dsm_example.png',
        'chapter': '5.2',
        'fig_num': '5-7',
        'caption': 'DSM高程图示例',
        'ref_text': '如图5-7展示了测试区域的DSM高程图，高程信息为PnP求解提供了关键的三维约束。',
    },
]


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=None, bold=None):
    """设置run的中英文字体和大小"""
    run.font.name = en_font
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def find_paragraph_index(doc, text_contains, style_name=None):
    """查找包含指定文本的段落索引"""
    for i, p in enumerate(doc.paragraphs):
        if text_contains in p.text:
            if style_name is None or p.style.name == style_name:
                return i
    return None


def find_paragraph_index_exact(doc, text_contains, start_from=0):
    """从指定位置开始查找包含指定文本的段落索引"""
    for i in range(start_from, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if text_contains in p.text:
            return i
    return None


def find_chapter_section_indices(doc, chapter_prefix):
    """找到章节的起始和结束段落索引（基于Heading 2）"""
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 2' and p.text.strip().startswith(chapter_prefix):
            start = i
        elif start is not None and p.style.name in ('Heading 1', 'Heading 2'):
            if not p.text.strip().startswith(chapter_prefix.split('.')[0] + '.'):
                end = i
                break
    if start is not None and end is None:
        # Find end of document
        end = len(doc.paragraphs)
    return start, end


def insert_paragraph_after(doc, ref_paragraph, text='', style=None):
    """在指定段落之后插入新段落"""
    new_p = doc.add_paragraph(text, style=style)
    # Move the new paragraph to after the reference paragraph
    ref_paragraph._element.addnext(new_p._element)
    return new_p


def insert_image_with_caption(doc, after_paragraph, img_filename, fig_num, caption_text, img_dir):
    """插入图片和图题"""
    img_path = os.path.join(img_dir, img_filename)
    if not os.path.exists(img_path):
        print(f'  WARNING: Image not found: {img_path}')
        return None, None

    # Add empty paragraph for spacing before image
    p_space_before = insert_paragraph_after(doc, after_paragraph, '')

    # Add image paragraph
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(14))
    after_paragraph._element.addnext(p_img._element)
    p_space_before._element.addnext(p_img._element)

    # Add caption paragraph (Heading 3 style)
    caption_full = f'图{fig_num} {caption_text}'
    p_caption = doc.add_paragraph(caption_full, style='Heading 3')
    p_img._element.addnext(p_caption._element)

    # Format caption
    for run in p_caption.runs:
        set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=10.5)

    # Add empty paragraph after caption for spacing
    p_space_after = doc.add_paragraph('')
    p_caption._element.addnext(p_space_after._element)

    return p_caption, p_space_after


def main():
    print(f'Loading document: {INPUT_FILE}')
    doc = Document(INPUT_FILE)
    print(f'  Paragraphs: {len(doc.paragraphs)}')
    print(f'  Sections: {len(doc.sections)}')
    print(f'  Tables: {len(doc.tables)}')

    # ═══════════════════════════════════════════
    # F8: 添加分节符
    # ═══════════════════════════════════════════
    print('\n=== F8: Adding section breaks ===')

    # 1. 封面后（"摘  要"前）
    abstract_idx = find_paragraph_index(doc, '摘  要')
    if abstract_idx is not None:
        print(f'  Found abstract at paragraph {abstract_idx}')
        # Insert section break before abstract
        # We need to add a section break to the paragraph just before "摘要"
        # The simplest approach: add section break to the paragraph before abstract
        prev_p = doc.paragraphs[abstract_idx - 1]
        # Add section break by setting the paragraph's section start type
        # Actually, we need to add a new section. Let's use a different approach:
        # We'll add the section break after the cover page content
        # Find the last cover paragraph (before abstract)
        last_cover_idx = abstract_idx - 1
        while last_cover_idx > 0 and not doc.paragraphs[last_cover_idx].text.strip():
            last_cover_idx -= 1

        # Add section break after the last cover paragraph
        # We need to insert a section break using XML manipulation
        from docx.oxml import OxmlElement
        # Create a new paragraph with section break
        p_break = doc.add_paragraph()
        # Move it before abstract
        doc.paragraphs[abstract_idx]._element.addprevious(p_break._element)
        # Add section break properties
        pPr = p_break._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_break._element.insert(0, pPr)
        sectPr = OxmlElement('w:sectPr')
        # Set section type to NEW_PAGE
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'continuous')
        sectPr.append(sectType)
        pPr.append(sectPr)
        print('  Added section break after cover page')

    # 2. 摘要后（"目  录"或TOC前）
    toc_idx = find_paragraph_index(doc, '【请右键更新域以生成目录】')
    if toc_idx is not None:
        print(f'  Found TOC placeholder at paragraph {toc_idx}')
        # Add section break before TOC
        p_break = doc.add_paragraph()
        doc.paragraphs[toc_idx]._element.addprevious(p_break._element)
        pPr = p_break._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_break._element.insert(0, pPr)
        sectPr = OxmlElement('w:sectPr')
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'continuous')
        sectPr.append(sectType)
        pPr.append(sectPr)
        print('  Added section break after abstract')

    # 3. 目录后（"第一章"前）
    first_chapter_idx = find_paragraph_index(doc, '第一章 研究背景与任务需求', 'Heading 1')
    if first_chapter_idx is not None:
        print(f'  Found first chapter at paragraph {first_chapter_idx}')
        p_break = doc.add_paragraph()
        doc.paragraphs[first_chapter_idx]._element.addprevious(p_break._element)
        pPr = p_break._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_break._element.insert(0, pPr)
        sectPr = OxmlElement('w:sectPr')
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'continuous')
        sectPr.append(sectType)
        pPr.append(sectPr)
        print('  Added section break after TOC')

    # 4. 参考文献前
    ref_idx = find_paragraph_index(doc, '参考文献', 'Heading 1')
    if ref_idx is not None:
        print(f'  Found references at paragraph {ref_idx}')
        p_break = doc.add_paragraph()
        doc.paragraphs[ref_idx]._element.addprevious(p_break._element)
        pPr = p_break._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_break._element.insert(0, pPr)
        sectPr = OxmlElement('w:sectPr')
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'continuous')
        sectPr.append(sectType)
        pPr.append(sectPr)
        print('  Added section break before references')

    # ═══════════════════════════════════════════
    # F9: 封面对齐（改为左对齐）
    # ═══════════════════════════════════════════
    print('\n=== F9: Cover page alignment ===')
    # Cover page paragraphs are indices 0-17 (before abstract)
    cover_keywords = ['项目名称', '项目负责人', '编撰', '审校', '北京市航空', '项目结题', '水面目标精准定位', '技术报告', '附件名称']
    for i in range(0, abstract_idx if abstract_idx else 18):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(f'  Set paragraph {i} to LEFT: {text[:40]}')

    # ═══════════════════════════════════════════
    # F10: 封面字体统一（宋体+Times New Roman）
    # ═══════════════════════════════════════════
    print('\n=== F10: Cover page fonts ===')
    for i in range(0, abstract_idx if abstract_idx else 18):
        p = doc.paragraphs[i]
        for run in p.runs:
            set_run_font(run, cn_font='宋体', en_font='Times New Roman')
        if p.text.strip():
            print(f'  Set font for paragraph {i}: {p.text.strip()[:40]}')

    # ═══════════════════════════════════════════
    # F11: 页眉页脚距边界
    # ═══════════════════════════════════════════
    print('\n=== F11: Header/footer distances ===')
    for i, sec in enumerate(doc.sections):
        sec.header_distance = Cm(1.42)
        sec.footer_distance = Cm(1.67)
        print(f'  Section {i}: header_distance=1.42cm, footer_distance=1.67cm')

    # ═══════════════════════════════════════════
    # F12: 删除"生成说明"等非正式内容
    # ═══════════════════════════════════════════
    print('\n=== F12: Remove generation notes ===')
    # The non-formal content starts from paragraph 305 ("局限性：") to the end
    # We need to keep the references (293-302) but remove everything after 302
    # Let's identify the content to remove
    # References end at paragraph 302
    # Non-formal content starts at 304 (empty) with "局限性：" at 305

    # Find the last reference paragraph
    last_ref_idx = None
    for i in range(ref_idx if ref_idx else 293, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text.startswith('[') and ']' in text:
            last_ref_idx = i

    if last_ref_idx:
        print(f'  Last reference at paragraph {last_ref_idx}')
        # Remove everything after last reference
        # We'll remove paragraphs from last_ref_idx+1 to end
        paragraphs_to_remove = []
        for i in range(last_ref_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if text:  # Only non-empty paragraphs
                paragraphs_to_remove.append(i)
                print(f'  Removing paragraph {i}: {text[:50]}')

        # Remove paragraphs by removing their XML elements
        for idx in reversed(paragraphs_to_remove):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)

    # Also check for empty paragraphs after references
    # Remove trailing empty paragraphs
    while len(doc.paragraphs) > 0:
        last_p = doc.paragraphs[-1]
        if not last_p.text.strip() and not last_p._element.findall(qn('w:r')):
            last_p._element.getparent().remove(last_p._element)
        else:
            break

    print(f'  Document now has {len(doc.paragraphs)} paragraphs')

    # ═══════════════════════════════════════════
    # F14: 参考文献字体统一为12pt
    # ═══════════════════════════════════════════
    print('\n=== F14: Reference font size ===')
    # Re-find references heading after potential removals
    ref_idx_new = find_paragraph_index(doc, '参考文献', 'Heading 1')
    if ref_idx_new is not None:
        print(f'  References heading at paragraph {ref_idx_new}')
        # Format all reference paragraphs
        for i in range(ref_idx_new + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = p.text.strip()
            if not text:
                continue
            # Stop if we hit non-reference content
            if text.startswith('局限性') or text.startswith('代码修改'):
                break
            for run in p.runs:
                run.font.size = Pt(12)
                set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12)
            print(f'  Set 12pt for paragraph {i}: {text[:50]}')

    # ═══════════════════════════════════════════
    # F15: Heading 2 字体统一
    # ═══════════════════════════════════════════
    print('\n=== F15: Heading 2 font unification ===')
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 2':
            for run in p.runs:
                set_run_font(run, cn_font='宋体', en_font='Times New Roman')
            print(f'  Unified font for Heading 2 at {i}: {p.text.strip()[:50]}')

    # ═══════════════════════════════════════════
    # 任务B: 插入图表
    # ═══════════════════════════════════════════
    print('\n=== Task B: Insert charts ===')

    # We need to insert charts after the relevant section content
    # Process charts in reverse order of their chapter position to avoid index shifting
    # First, map chapter prefixes to their paragraph indices

    chapter_map = {}
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 2':
            text = p.text.strip()
            # Extract chapter number (e.g., "5.5" from "5.5 视觉定位算法结果分析")
            parts = text.split()
            if parts and '.' in parts[0]:
                chapter_map[parts[0]] = i

    print(f'  Found chapter headings: {list(chapter_map.keys())}')

    # Process charts from last chapter to first to avoid index shifting
    sorted_charts = sorted(CHARTS, key=lambda c: c['chapter'], reverse=True)

    for chart in sorted_charts:
        chapter = chart['chapter']
        fig_num = chart['fig_num']
        caption = chart['caption']
        img_file = chart['file']
        ref_text = chart['ref_text']

        print(f'\n  Processing {fig_num}: {caption} (chapter {chapter})')

        # Find the chapter heading
        chapter_idx = chapter_map.get(chapter)
        if chapter_idx is None:
            print(f'    WARNING: Chapter {chapter} not found, skipping')
            continue

        # Find the next chapter heading (to know where this chapter ends)
        next_chapter_idx = None
        for i in range(chapter_idx + 1, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            if p.style.name in ('Heading 1', 'Heading 2'):
                next_chapter_idx = i
                break
        if next_chapter_idx is None:
            next_chapter_idx = len(doc.paragraphs)

        # Find a good insertion point: after the last content paragraph before next chapter
        # We'll insert after the last non-empty paragraph before next chapter
        insert_after_idx = next_chapter_idx - 1
        while insert_after_idx > chapter_idx:
            if doc.paragraphs[insert_after_idx].text.strip():
                break
            insert_after_idx -= 1

        # If we found a table in the section, insert after the table
        # Check for tables between chapter_idx and next_chapter_idx
        # We need to find the last paragraph that's followed by a table or is content
        # Let's just insert after the last content paragraph

        insert_after_p = doc.paragraphs[insert_after_idx]
        print(f'    Inserting after paragraph {insert_after_idx}: {insert_after_p.text.strip()[:50]}')

        # Add reference text to the content
        p_ref = insert_paragraph_after(doc, insert_after_p, ref_text)
        # Format the reference text
        for run in p_ref.runs:
            set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12)

        # Insert image with caption
        img_path = os.path.join(IMG_DIR, img_file)
        if not os.path.exists(img_path):
            print(f'    WARNING: Image not found: {img_path}')
            continue

        # Add image paragraph
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(img_path, width=Cm(14))
        p_ref._element.addnext(p_img._element)

        # Add caption
        caption_full = f'图{fig_num} {caption}'
        p_caption = doc.add_paragraph(caption_full, style='Heading 3')
        p_img._element.addnext(p_caption._element)

        # Format caption
        for run in p_caption.runs:
            set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=10.5)

        print(f'    Inserted image: {img_file}')
        print(f'    Inserted caption: {caption_full}')

    # ═══════════════════════════════════════════
    # 保存文档
    # ═══════════════════════════════════════════
    print(f'\n=== Saving document: {OUTPUT_FILE} ===')
    doc.save(OUTPUT_FILE)
    print(f'Done! Saved to: {OUTPUT_FILE}')
    print(f'Final paragraph count: {len(doc.paragraphs)}')


if __name__ == '__main__':
    main()
