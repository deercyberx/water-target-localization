# -*- coding: utf-8 -*-
"""
修正版：插入图片到正确位置
- 1_Retrieval.png → 5.5节（图5-11）
- 1-1581.png → 5.10节（图5-12）
"""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE = '水面目标精准定位技术报告_终稿_v4.docx'
OUTPUT_FILE = '水面目标精准定位技术报告_终稿_v5.docx'

IMG_DIR = '素材/截图/G10_matching_QZ_SongCity_1'

def insert_image_with_caption(doc, para_index, image_path, ref_text, caption_text, width_inches=5.33):
    """在指定段落后插入图片、引用文本和图题"""
    target_para = doc.paragraphs[para_index]

    # 1. 插入引用文本段落
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref_run = ref_para.add_run(ref_text)
    ref_run.font.size = Pt(12)
    target_para._element.addnext(ref_para._element)

    # 2. 插入图片段落
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    ref_para._element.addnext(img_para._element)

    # 3. 插入图题段落
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.style = doc.styles['Heading 3']
    caption_run = caption_para.add_run(caption_text)
    caption_run.font.size = Pt(10.5)
    caption_run.bold = True
    img_para._element.addnext(caption_para._element)

    return caption_para

def main():
    print(f'读取: {INPUT_FILE}')
    doc = Document(INPUT_FILE)

    print('=' * 50)
    print('插入最终图片（修正版）')
    print('=' * 50)

    # 1. 插入 1_Retrieval.png 到 5.5节
    print('\n[1] 插入检索结果可视化（图5-11）')
    print('  位置: 5.5节，图5-2后面')

    # 找到"图5-2 定位误差累积分布函数"的Heading 3段落
    idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 3' and '图5-2' in p.text:
            idx = i
            break

    if idx >= 0:
        # 找到图题后面的Normal段落
        for i in range(idx+1, min(idx+5, len(doc.paragraphs))):
            if doc.paragraphs[i].style.name == 'Normal' and doc.paragraphs[i].text.strip():
                idx = i
                break

        ref_text = '如图5-11所示，以QZ_SongCity场景为例，CAMP检索模型返回的Top-5结果中，前两名与查询图像具有较高的空间重叠度（PDE分别为0.109和0.603），而第三名开始出现显著偏差（PDE=1.603）。这说明检索质量直接决定了后续匹配的成功率。'
        caption = '图5-11 检索结果Top5可视化示例（QZ_SongCity场景，CAMP检索模型）'

        insert_image_with_caption(doc, idx, f'{IMG_DIR}/1_Retrieval.png', ref_text, caption)
        print(f'  ✅ 插入到段落 {idx} 后面')
    else:
        print('  ❌ 未找到图5-2位置')

    # 2. 插入 1-1581.png 到 5.10节
    print('\n[2] 插入匹配点对可视化（图5-12）')
    print('  位置: 5.10节，段落298后面')

    # 找到5.10节的最后一个段落（段落298）
    idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 2' and '5.10' in p.text:
            # 找到5.10节的Heading 2，然后找该节最后一个Normal段落
            for j in range(i+1, min(i+15, len(doc.paragraphs))):
                if doc.paragraphs[j].style.name == 'Heading 2':
                    # 到了下一个Heading 2，取前一个Normal段落
                    for k in range(j-1, i, -1):
                        if doc.paragraphs[k].style.name == 'Normal' and doc.paragraphs[k].text.strip():
                            idx = k
                            break
                    break
            break

    if idx >= 0:
        ref_text = '如图5-12所示，当检索Top-1结果与真实位置高度重合时（PDE=0.109），RoMa密集匹配可产生1581个内点，为PnP求解提供充足的几何约束。匹配点对覆盖了图像中的建筑物、道路等刚性特征，有效滤除了水面等弱纹理区域的干扰。'
        caption = '图5-12 匹配点对可视化（Top-1检索结果，1581内点，PDE=0.109）'

        insert_image_with_caption(doc, idx, f'{IMG_DIR}/1-1581.png', ref_text, caption)
        print(f'  ✅ 插入到段落 {idx} 后面')
    else:
        print('  ❌ 未找到5.10节位置')

    # 保存
    print(f'\n保存: {OUTPUT_FILE}')
    doc.save(OUTPUT_FILE)
    print('完成!')

    # 统计
    doc2 = Document(OUTPUT_FILE)
    print(f'\n最终统计:')
    print(f'  段落数: {len(doc2.paragraphs)}')
    print(f'  表格数: {len(doc2.tables)}')
    print(f'  图片数: {len(doc2.inline_shapes)}')

if __name__ == '__main__':
    main()
