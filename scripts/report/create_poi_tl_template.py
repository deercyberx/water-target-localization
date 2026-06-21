# -*- coding: utf-8 -*-
"""
创建poi-tl模板：为每个表格预定义表头，数据行用占位符
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 样式
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = 'SimSun'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level, size in [(1, 16), (3, 12)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.name = 'SimHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

def add_poi_table(doc, caption, headers, data_keys):
    """创建poi-tl表格：表头固定，数据行用{{key}}占位"""
    doc.add_paragraph(caption, style='Heading 3')
    table = doc.add_table(rows=2, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i, k in enumerate(data_keys):
        table.rows[1].cells[i].text = '{{' + k + '}}'

# ===== 第五章 =====
doc.add_heading('第五章 实验验证与结果分析', level=1)

# 5.1
doc.add_heading('5.1 实验数据与测试场景', level=1)
doc.add_paragraph('{{s5_1_intro}}')
add_poi_table(doc, '表5-1 AnyVisLoc数据集属性',
    ['属性', '内容'], ['t1_col1', 't1_col2'])
add_poi_table(doc, '表5-2 Demo测试子集说明',
    ['测试区域', '查询图像数', '航空参考图', '卫星参考图', '姿态元数据'],
    ['t2_col1', 't2_col2', 't2_col3', 't2_col4', 't2_col5'])
doc.add_paragraph('图5-1 测试区域图像样例', style='Heading 3')
doc.add_paragraph('{{@fig5_1}}')
doc.add_paragraph('{{s5_1_note}}')

# 5.2
doc.add_heading('5.2 参考地图与数字正射影像图说明', level=1)
doc.add_paragraph('{{s5_2_intro}}')
add_poi_table(doc, '表5-3 参考地图与高程数据说明',
    ['参考数据类型', '主要内容', '链路作用', '分辨率', '注意事项'],
    ['t3_col1', 't3_col2', 't3_col3', 't3_col4', 't3_col5'])
doc.add_paragraph('{{s5_2_note}}')

# 5.3
doc.add_heading('5.3 实验平台与参数设置', level=1)
doc.add_paragraph('{{s5_3_intro}}')
add_poi_table(doc, '表5-4 实验平台与运行环境',
    ['项目', '配置'], ['t4_col1', 't4_col2'])
add_poi_table(doc, '表5-5 定位链路关键参数设置',
    ['参数', '值', '影响说明'], ['t5_col1', 't5_col2', 't5_col3'])
doc.add_paragraph('{{s5_3_note}}')

# 5.4
doc.add_heading('5.4 评价指标', level=1)
doc.add_paragraph('{{s5_4_intro}}')
add_poi_table(doc, '表5-6 评价指标说明',
    ['指标名称', '指标含义', '评价方向', '适用实验', '对应链路环节'],
    ['t6_col1', 't6_col2', 't6_col3', 't6_col4', 't6_col5'])
doc.add_paragraph('{{s5_4_note}}')

# 5.5
doc.add_heading('5.5 视觉定位算法结果分析', level=1)
doc.add_paragraph('{{s5_5_intro}}')
add_poi_table(doc, '表5-7 视觉定位算法主结果',
    ['参考图类型', '样本数', 'A@5m', 'A@10m', 'A@20m', '平均误差', '中位数误差'],
    ['t7_col1', 't7_col2', 't7_col3', 't7_col4', 't7_col5', 't7_col6', 't7_col7'])
doc.add_paragraph('{{s5_5_analysis}}')
doc.add_paragraph('图5-3 视觉定位算法完整流程可视化', style='Heading 3')
doc.add_paragraph('{{@fig5_3}}')
doc.add_paragraph('{{s5_5_note}}')

# 5.6
doc.add_heading('5.6 不同参考地图条件下的定位效果分析', level=1)
doc.add_paragraph('{{s5_6_intro}}')
add_poi_table(doc, '表5-8 不同参考地图条件下的定位效果',
    ['参考图类型', '影像分辨率', 'DSM分辨率', 'A@5m', 'A@10m', 'A@20m', '平均误差'],
    ['t8_col1', 't8_col2', 't8_col3', 't8_col4', 't8_col5', 't8_col6', 't8_col7'])
doc.add_paragraph('图5-4 航空图与卫星图定位效果对比', style='Heading 3')
doc.add_paragraph('{{@fig5_4}}')
doc.add_paragraph('{{s5_6_note}}')

# 5.7
doc.add_heading('5.7 不同飞行高度和视角条件下的定位效果分析', level=1)
doc.add_paragraph('{{s5_7_intro}}')
add_poi_table(doc, '表5-9 不同高度与视角条件下定位效果',
    ['分析维度', '分组范围', '样本数', 'A@5m', 'A@20m', '平均误差', '解释重点'],
    ['t9_col1', 't9_col2', 't9_col3', 't9_col4', 't9_col5', 't9_col6', 't9_col7'])
doc.add_paragraph('图5-5 不同高度与视角条件下定位效果分析', style='Heading 3')
doc.add_paragraph('{{@fig5_5}}')
doc.add_paragraph('{{s5_7_note}}')

# 5.8
doc.add_heading('5.8 定位策略与匹配方法对比分析', level=1)
doc.add_paragraph('{{s5_8_intro}}')
add_poi_table(doc, '表5-10 定位策略、匹配方法与先验噪声对比',
    ['分析对象', '对比设置', 'A@5m', 'A@10m', 'A@20m', '主要结论'],
    ['t10_col1', 't10_col2', 't10_col3', 't10_col4', 't10_col5', 't10_col6'])
doc.add_paragraph('图5-6 定位策略、匹配方法与先验噪声对比分析', style='Heading 3')
doc.add_paragraph('{{@fig5_6}}')
doc.add_paragraph('{{s5_8_note}}')

# 5.9
doc.add_heading('5.9 与目标检测结果联动的定位误差分析框架', level=1)
doc.add_paragraph('{{s5_9_intro}}')
add_poi_table(doc, '表5-11 检测框联动定位误差分析框架',
    ['检测框输入类型', '扰动设置', '所需参数', '输出指标', '需回答的问题'],
    ['t11_col1', 't11_col2', 't11_col3', 't11_col4', 't11_col5'])
doc.add_paragraph('图5-7 检测框→定位误差传播链路', style='Heading 3')
doc.add_paragraph('{{@fig5_7}}')
doc.add_paragraph('{{s5_9_note}}')

# 5.10
doc.add_heading('5.10 典型定位结果与失败案例分析', level=1)
doc.add_paragraph('{{s5_10_intro}}')
add_poi_table(doc, '表5-12 典型案例选择与失败原因说明',
    ['案例类型', '场景特征', '结果现象', '主要原因', '正文分析重点'],
    ['t12_col1', 't12_col2', 't12_col3', 't12_col4', 't12_col5'])
doc.add_paragraph('图5-8 典型定位结果与失败案例分析', style='Heading 3')
doc.add_paragraph('{{@fig5_8}}')
doc.add_paragraph('{{s5_10_note}}')

# 5.11
doc.add_heading('5.11 本章小结', level=1)
doc.add_paragraph('{{s5_11_summary}}')
doc.add_paragraph('{{s5_11_c1}}')
doc.add_paragraph('{{s5_11_c2}}')
doc.add_paragraph('{{s5_11_c3}}')
doc.add_paragraph('{{s5_11_c4}}')
doc.add_paragraph('{{s5_11_c5}}')
doc.add_paragraph('{{s5_11_c6}}')

output = r'C:\Users\deer\Desktop\Projects\poi-tl-project\poi-tl\templates\chapter5_template.docx'
doc.save(output)
print(f'Template saved: {output}')
