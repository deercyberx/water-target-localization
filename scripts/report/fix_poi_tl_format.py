# -*- coding: utf-8 -*-
"""
修复poi-tl生成的docx格式问题：
1. 表头背景色
2. 表头粗体
3. 表格数据居中
4. 图片尺寸缩放
5. 图片居中
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading_elm.append(shading)

def fix_tables(doc):
    """修复表格格式：表头背景色+粗体，数据居中"""
    for table in doc.tables:
        # 表头行
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                # P1: 背景色
                set_cell_shading(cell, 'D9E2F3')
                # P2: 粗体
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

        # P3: 所有数据居中
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print(f"  表格修复: {len(doc.tables)}个")

def fix_images(doc):
    """修复图片尺寸和居中"""
    max_width = Cm(15)  # 最大宽度15cm

    for para in doc.paragraphs:
        has_img = False
        for run in para.runs:
            # 检查是否有图片
            drawings = run._element.findall(qn('w:drawing'))
            if drawings:
                has_img = True
                for drawing in drawings:
                    # 查找wp:inline或wp:anchor
                    for inline in drawing.findall('.//' + qn('wp:inline')):
                        extent = inline.find(qn('wp:extent'))
                        if extent is not None:
                            cx = int(extent.get('cx', 0))
                            cy = int(extent.get('cy', 0))

                            # 缩放
                            if cx > max_width:
                                ratio = max_width / cx
                                new_cx = int(max_width)
                                new_cy = int(cy * ratio)
                                extent.set('cx', str(new_cx))
                                extent.set('cy', str(new_cy))

                                # 更新图片内部尺寸
                                for ext in drawing.findall('.//' + qn('a:ext')):
                                    ext.set('cx', str(new_cx))
                                    ext.set('cy', str(new_cy))

        # P5: 图片段落居中
        if has_img:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print("  图片修复: 完成")

def main():
    input_path = r'C:\Users\deer\Desktop\Projects\Research\projects\水上目标定位\reports\第五章_实验验证与结果分析_poi-tl_final.docx'
    output_path = r'C:\Users\deer\Desktop\Projects\Research\projects\水上目标定位\reports\第五章_实验验证与结果分析_poi-tl_v2.docx'

    print("读取文档...")
    doc = Document(input_path)

    print("修复表格...")
    fix_tables(doc)

    print("修复图片...")
    fix_images(doc)

    print("保存...")
    doc.save(output_path)
    print(f"完成: {output_path}")

if __name__ == '__main__':
    main()
